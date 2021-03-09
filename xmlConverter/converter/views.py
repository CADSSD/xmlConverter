from django.shortcuts import render, redirect
from django.http import HttpResponse
from datetime import datetime
from django.templatetags.static import static

# funtional
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

import pdfkit
from htmlBuilder.tags import *
from htmlBuilder.attributes import Class, Style as InlineStyle

# Create your views here.

def upload(request):
  return render(request, 'upload/upload.html')

def message(request):
  return HttpResponse("Caricamento completato")

def convert_in_word(request):
  document = Document()
  style = document.styles['Normal']
  font = style.font
  font.name = 'Calibri'
  font.size = Pt(10)

  margin = 1.5
  sections = document.sections
  for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)
  
  footer_section = sections[0]
  footer = footer_section.footer

  tree = ET.parse(request.FILES.get('file'))
  root = tree.getroot()
  
  exporter_ragione_sociale = root.find("./Esportatore_1/RagioneSociale").text
  exporter_address = root.find("./Esportatore_1/Indirizzo").text
  exporter_comune = root.find("./Esportatore_1/Comune").text
  exporter_cap = root.find("./Esportatore_1/CAP").text
  exporter_country_code = root.find("./Esportatore_1/CodPaese").text
  exporter_string = exporter_ragione_sociale + "\n" + exporter_address + "\n" + exporter_comune + " " + exporter_cap + "\n" + exporter_country_code

  provenience = root.find("./Origine_4").text

  destinatario_ragione_sociale = root.find("./Destinatario_3/RagioneSociale").text
  destinatario_address = root.find("./Destinatario_3/Indirizzo").text
  destinatario_comune = root.find("./Destinatario_3/Comune").text
  destinatario_cap = root.find("./Destinatario_3/CAP").text
  destinatario_country_code = root.find("./Destinatario_3/CodPaese").text
  destinatario_string = destinatario_ragione_sociale + "\n" + destinatario_address + "\n" + destinatario_comune + " " + destinatario_cap + "\n" + destinatario_country_code

  arrival = root.find("./Destinazione_5").text

  table = document.add_table(rows=0, cols=2)
  row_cells = table.add_row().cells
  row_cells[0].text = exporter_string
  row_cells[0].style = document.styles['Normal']

  par = row_cells[1].add_paragraph()
  par.text = provenience
  par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  par.style = document.styles['Normal']

  row_cells = table.add_row().cells
  row_cells[0].text = destinatario_string
  row_cells[0].style = document.styles['Normal']

  par = row_cells[1].add_paragraph()
  par.text = arrival
  par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  par.style = document.styles['Normal']

  table = document.add_table(rows=0, cols=2)
  row_cells = table.add_row().cells
  nested_table = row_cells[1].add_table(rows=0, cols=2)
  nested_table_cells = nested_table.add_row().cells
  nested_table_cells[0].text = provenience
  nested_table_cells[1].text = arrival

  document.add_paragraph('\n\n\n\n\n\n\n\n\n\n')

  table = document.add_table(rows=0, cols=2)
  for casella_8 in root.findall("./Casella_8_9_10/Casella_8"):
    row_cells = table.add_row().cells
    row_cells[0].text = casella_8.find("./Progressivo").text + " " + casella_8.find("./TotaleColli").text + casella_8.find("./CodiceConfezione").text + casella_8.find("./DescrizioneMerce").text
    row_cells[0].style = document.styles['Normal']
    row_cells[1].text = casella_8.find("./Progressivo").text
    row_cells[1].style = document.styles['Normal']


  cert_id = root.find("./Cert_ID").text
  visto_modello = root.find("./VistoDogana_11/Modello").text
  visto_numero = root.find("./VistoDogana_11/Numero").text
  luogo = root.find("./VistoDogana_11/Luogo").text
  data = root.find("./VistoDogana_11/Data").text[0: 10]
  data2 = root.find("./VistoDogana_11/Del").text[0: 10]

  data = datetime.strptime(data, "%Y-%m-%d")
  data = datetime.strftime(data, "%d/%m/%Y")

  data2 = datetime.strptime(data2, "%Y-%m-%d")
  data2 = datetime.strftime(data2, "%d/%m/%Y")

  footer_string = "Certificato" + cert_id + "\n" + "Versione 1\n\n\n"
  footer_string += visto_modello + "     " + visto_numero + "\t\t" + luogo + "  " + data + "\n"
  footer_string += data2 + "\n" + "28100 UD PARMA"
  footer_par = footer.paragraphs[0]
  footer_par.text = footer_string

  # document.add_page_break()
  dirspot = os.getcwd()
  dateTimeObj = datetime.now()
  timestampStr = dateTimeObj.strftime("%d%b%Y%H%M%S%f")
  document.save(dirspot + '/converter/static/files/' + timestampStr + '.docx')
  url = static('files/' + timestampStr + '.docx')

  return HttpResponse(url)

  # return HttpResponse("Caricamento completato")

def convert_in_pdf(request):
  tree = ET.parse(request.FILES.get('file'))
  root = tree.getroot()

  exporter_ragione_sociale = root.find("./Esportatore_1/RagioneSociale").text
  exporter_address = root.find("./Esportatore_1/Indirizzo").text
  exporter_comune = root.find("./Esportatore_1/Comune").text
  exporter_cap = root.find("./Esportatore_1/CAP").text
  exporter_country_code = root.find("./Esportatore_1/CodPaese").text
  exporter_string = exporter_ragione_sociale + "\n" + exporter_address + "\n" + exporter_comune + " " + exporter_cap + "\n" + exporter_country_code

  provenience = root.find("./Origine_4").text

  destinatario_ragione_sociale = root.find("./Destinatario_3/RagioneSociale").text
  destinatario_address = root.find("./Destinatario_3/Indirizzo").text
  destinatario_comune = root.find("./Destinatario_3/Comune").text
  destinatario_cap = root.find("./Destinatario_3/CAP").text
  destinatario_country_code = root.find("./Destinatario_3/CodPaese").text
  destinatario_string = destinatario_ragione_sociale + "\n" + destinatario_address + "\n" + destinatario_comune + " " + destinatario_cap + "\n" + destinatario_country_code

  arrival = root.find("./Destinazione_5").text

  cert_id = root.find("./Cert_ID").text
  visto_modello = root.find("./VistoDogana_11/Modello").text
  visto_numero = root.find("./VistoDogana_11/Numero").text
  luogo = root.find("./VistoDogana_11/Luogo").text
  data = root.find("./VistoDogana_11/Data").text[0: 10]
  data2 = root.find("./VistoDogana_11/Del").text[0: 10]
  data = datetime.strptime(data, "%Y-%m-%d")
  data = datetime.strftime(data, "%d/%m/%Y")
  data2 = datetime.strptime(data2, "%Y-%m-%d")
  data2 = datetime.strftime(data2, "%d/%m/%Y")

  dirspot = os.getcwd()
  dateTimeObj = datetime.now()
  timestampStr = dateTimeObj.strftime("%d%b%Y%H%M%S%f")

  html = Html([],
    Head([],
      Title([], timestampStr)
    ),
    Body([InlineStyle(font_size="10pt")],
      Div([InlineStyle(width="fit-content", position="absolute", top="70px", left="80px")],
        exporter_ragione_sociale,
        Br([]),
        exporter_address,
        Br([]),
        exporter_comune + " " + exporter_cap,
        Br([]),
        exporter_country_code
      ),
      Div([InlineStyle(width="fit-content", position="absolute", top="163px", left="680px")],
        provenience
      ),
      Div([InlineStyle(width="fit-content", position="absolute", top="200px", left="80px")],
        destinatario_ragione_sociale,
        Br([]),
        destinatario_address,
        Br([]),
        destinatario_comune + " " + destinatario_cap,
        Br([]),
        destinatario_country_code
      ),
      Div([InlineStyle(width="fit-content", position="absolute", top="238px", left="680px")],
        arrival
      ),
      Div([InlineStyle(width="fit-content", position="absolute", top="360px", left="570px")],
        provenience
      ),
      Div([InlineStyle(width="fit-content", position="absolute", top="360px", left="745px")],
        arrival
      ),
      Div([InlineStyle(width="500px", position="absolute", top="715px", left="55px")],
        [Div([],
          casella_8.find("./Progressivo").text + " ",
          Span([InlineStyle(color="white")], casella_8.find("./TotaleColli").text + casella_8.find("./CodiceConfezione").text),
          casella_8.find("./DescrizioneMerce").text,
          Span([InlineStyle(position="relative", bottom="50px", left="-20px")], casella_8.find("./TotaleColli").text + casella_8.find("./CodiceConfezione").text)
        ) for casella_8 in root.findall("./Casella_8_9_10/Casella_8")]
      ),
      Div([InlineStyle(width="fit-content", position="absolute", top="715px", left="730px")],
        Div([],
          root.find("./Casella_8_9_10/Casella_9/PesoLordo").text
        )
      ),
      Div([InlineStyle(width="fit-content", position="absolute", top="715px", left="840px")],
        Div([],
          [
            Span([], 
              fattura.text + " ",
              Br([])
            ) for fattura in root.find('./Casella_8_9_10/Casella_10').findall("./NumeroFattura")]
        )
      ),
      Div([InlineStyle(position="absolute", bottom="232px", left="230")],
        "Certificato " + cert_id,
        Br(),
        "Versione 1"
      ),
      Div([InlineStyle(position="absolute", bottom="134px", left="158")],
        visto_modello + "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;" + visto_numero
      ),
      Div([InlineStyle(position="absolute", bottom="115px", left="178")],
        data2,
      ),
      Div([InlineStyle(position="absolute", bottom="88px", left="178")],
        Br([]),
        Br([]),
        "28100 UD PARMA"
      ),
      Div([InlineStyle(position="absolute", bottom="70px", left="736")],
        luogo + "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;" + data
      ),
      Div([InlineStyle(position="absolute", bottom="29px", left="138")],
        luogo + "&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;" + data
      ),
    )
  )

  print(html.render(pretty=True))

  options = {
    'page-size': 'A4',
    'no-outline': None
  }

  pdfkit.from_string(html.render(), dirspot + '/converter/static/files/' + timestampStr + '.pdf', options)

  url = static('files/' + timestampStr + '.pdf')

  return HttpResponse(url)